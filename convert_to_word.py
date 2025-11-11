#!/usr/bin/env python3
"""
Script para converter MANUAL_DO_USUARIO.md para Word (.docx)
Requisitos: pip install python-docx
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Erro: Biblioteca python-docx não encontrada.")
    print("\n📦 Instale com: pip install python-docx")
    sys.exit(1)

def process_text_with_formatting(paragraph, text):
    """Processa texto com formatação markdown (negrito, itálico, código)"""
    # Processar negrito (**texto**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

def parse_markdown_to_docx(md_file, docx_file):
    """Converte arquivo Markdown para Word"""
    
    # Ler o arquivo Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Criar documento Word
    doc = Document()
    
    # Configurar estilo padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Dividir em linhas
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Título nível 1 (#)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        # Título nível 2 (##)
        if line.startswith('## '):
            title = line[3:].strip()
            heading = doc.add_heading(title, level=2)
            i += 1
            continue
        
        # Título nível 3 (###)
        if line.startswith('### '):
            title = line[4:].strip()
            heading = doc.add_heading(title, level=3)
            i += 1
            continue
        
        # Lista numerada
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Lista com marcadores (- ou *)
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            
            # Processar texto em negrito dentro da lista
            if '**' in text or '`' in text:
                process_text_with_formatting(p, text)
            else:
                p.add_run(text)
            
            i += 1
            continue
        
        # Código inline (`)
        if '`' in line:
            # Processar código inline
            parts = re.split(r'(`[^`]+`)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    p.add_run(part)
            i += 1
            continue
        
        # Bloco de código (```)
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            continue
        
        # Linha vazia
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        
        # Texto normal com formatação
        if line:
            # Processar texto com formatação
            if '**' in line or '`' in line:
                p = doc.add_paragraph()
                process_text_with_formatting(p, line)
            else:
                # Texto normal
                p = doc.add_paragraph(line)
            
            i += 1
    
    # Salvar documento
    doc.save(docx_file)
    print(f"✅ Documento Word criado com sucesso: {docx_file}")

if __name__ == "__main__":
    md_file = Path(__file__).parent / "MANUAL_DO_USUARIO.md"
    docx_file = Path(__file__).parent / "MANUAL_DO_USUARIO.docx"
    
    if not md_file.exists():
        print(f"❌ Erro: Arquivo {md_file} não encontrado!")
        sys.exit(1)
    
    print(f"📄 Convertendo {md_file} para Word...")
    parse_markdown_to_docx(md_file, docx_file)

